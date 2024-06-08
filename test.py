import time
import random


def generate_population(size, num_items):
    """Tạo một quần thể các giải pháp ngẫu nhiên."""
    return [[random.randint(0, 1) for _ in range(num_items)] for _ in range(size)]


def fitness(solution, weights, values, capacity):
    """Tính độ thích nghi của một giải pháp."""
    total_weight = sum(weight * item for weight, item in zip(weights, solution))
    total_value = sum(value * item for value, item in zip(values, solution))
    if total_weight > capacity:
        return 0  # Giải pháp không khả thi
    return total_value


def tournament_selection(population, fitness_values, tournament_size=2):
    selected = []
    # """Lựa chọn cá thể elite"""
    # sorted_population = sorted(population, key=lambda genome: fitness(genome, weights, values, capacity) , reverse=True)
    # selected.append(sorted_population[0])
    """Thực hiện lựa chọn giải đấu để chọn cha mẹ."""
    for _ in range(2):  # Chọn 2 cha mẹ
        tournament = random.sample(range(len(population)), tournament_size)
        best_index = max(tournament, key=lambda i: fitness_values[i])
        selected.append(population[best_index])
    return selected

def fitness_scaling(fitnesses, scaling_method="linear"):
    """Điều chỉnh giá trị fitness bằng phương pháp scaling."""
    if scaling_method == "linear":
        min_fitness = min(fitnesses)
        max_fitness = max(fitnesses)
        if max_fitness == min_fitness: # Tránh chia cho 0
            return [1 for _ in fitnesses]  # Nếu tất cả fitness bằng nhau, gán fitness scaled là 1
        scaled_fitnesses = [(f - min_fitness) / (max_fitness - min_fitness) for f in fitnesses]
    # Thêm các phương pháp scaling khác ở đây (ví dụ: sigma scaling, power law scaling)
    return scaled_fitnesses

def tournament_selection_with_scaling(population, fitnesses, tournament_size=2, scaling_method="linear"):
    """Lựa chọn cá thể bằng Tournament Selection với Fitness Scaling."""
    scaled_fitnesses = fitness_scaling(fitnesses, scaling_method)
    selected = []
    for _ in range(2):
        candidates = random.sample(population, tournament_size)
        # Chọn cá thể có fitness scaled cao nhất
        winner = max(candidates, key=lambda x: scaled_fitnesses[population.index(x)])
        selected.append(winner)
    return selected

    


        
### ---------CROSSOVER---------



def one_point_crossover(parent1, parent2):
    """Thực hiện lai ghép giữa hai giải pháp cha mẹ."""
    crossover_point = random.randint(1, len(parent1) - 1)
    child1 = parent1[:crossover_point] + parent2[crossover_point:]
    child2 = parent2[:crossover_point] + parent1[crossover_point:]
    return child1, child2


def two_points_crossover(parent1, parent2):
    """Lai ghép hai điểm."""
    # Chọn 2 điểm cắt ngẫu nhiên, đảm bảo điểm đầu nhỏ hơn điểm cuối
    point1 = random.randint(1, len(parent1) - 2)  
    point2 = random.randint(point1 + 1, len(parent1) - 1)
   
    # Tạo con bằng cách kết hợp các đoạn từ bố mẹ
    child1 = parent1[:point1] + parent2[point1:point2] + parent1[point2:]
    child2 = parent2[:point1] + parent1[point1:point2] + parent2[point2:]
    return child1, child2


 
def uniform_crossover(parent1, parent2):
    """Thực hiện lai ghép đồng nhất giữa hai giải pháp cha mẹ."""
    child1 = [parent1[i] if random.random() < 0.5 else parent2[i] for i in range(len(parent1))]
    child2 = [parent2[i] if random.random() < 0.5 else parent1[i] for i in range(len(parent1))]
    return child1, child2




### ---------MUTATION---------


def bit_flip(solution, mutation_rate):
    """Thực hiện đột biến trên một giải pháp."""
    for i in range(len(solution)):
        if random.random() < mutation_rate:
            solution[i] = 1 - solution[i]
    return solution


def swap_mutate(solution, mutation_rate):
    """Thực hiện đột biến swap trên một giải pháp."""
    for _ in range(len(solution)):
        if random.random() < mutation_rate:
            i, j = random.sample(range(len(solution)), 2)  # Chọn ngẫu nhiên 2 vị trí
            solution[i], solution[j] = solution[j], solution[i]  # Hoán đổi giá trị
    return solution


def inversion_mutate(solution, mutation_rate):
    """Thực hiện đột biến đảo đoạn trên một giải pháp."""
    if random.random() < mutation_rate:
        start = random.randint(0, len(solution) - 2)
        end = random.randint(start + 1, len(solution))
        solution[start:end] = reversed(solution[start:end])  # Đảo đoạn
    return solution


### ---------genetic_algorithm---------

def bruteforce_knapsack(weights, values, capacity):
    n = len(weights)  # Số lượng vật phẩm
    max_value = 0
    best_combination = []

    # Duyệt qua tất cả các tổ hợp có thể của vật phẩm (2^n tổ hợp)
    for i in range(2**n):
        combination = []
        total_weight = 0
        total_value = 0

        # Kiểm tra từng bit trong i để xác định vật phẩm nào được chọn
        for j in range(n):
            if (i >> j) & 1:  # Kiểm tra xem bit thứ j của i có bằng 1 không
                combination.append(j)
                total_weight += weights[j]
                total_value += values[j]

        # Kiểm tra xem tổ hợp có khả thi (không vượt quá sức chứa) và có giá trị lớn hơn không
        if total_weight <= capacity and total_value > max_value:
            max_value = total_value
            best_combination = combination

    return max_value, best_combination


def genetic_algorithm( result , crossover_func, mutation_func,weights, values, capacity, population_size, generations, mutation_rate ): 
#Giải bài toán Knapsack 0/1 bằng giải thuật di truyền 

    # Khởi tạo quần thể ban đầu 
    population = generate_population(population_size, len(weights)) 
    
    start = time.time() 
    for i in range(generations): 
        
        population = sorted(population, key=lambda genome: fitness(genome, weights, values, capacity) , reverse=True) 
        elit= fitness(population[0], weights, values, capacity)
        if ((elit >= result)) :
            
            return elit/result ,i,  (end - start)
        new_population =  population[0:2] 
        while len(new_population) < population_size: 
            fitness_values = [fitness(solution, weights, values, capacity) for solution in new_population] 
            parent1, parent2 = tournament_selection(new_population, fitness_values) 
            child1, child2 = crossover_func(parent1, parent2) 
            child1 = mutation_func(child1, mutation_rate) 
            child2 = mutation_func(child2, mutation_rate) 
            new_population.extend([child1, child2]) 

        population = new_population 
        end = time.time()
        

    best_solution = max(population, key=lambda x: fitness(x, weights, values, capacity)) 
    best_fitness = fitness(best_solution, weights, values, capacity) 
    return best_fitness/ result,i ,(end - start)


### ---------run_all_combinations---------


def run_all_combinations(result,weights, values, capacity, population_size=100, generations=100, mutation_rate=0.2):
    """Chạy thuật toán di truyền với tất cả các kết hợp của các hàm lai ghép và đột biến."""
    results = {}


    crossover_funcs = [one_point_crossover, two_points_crossover, uniform_crossover]
    mutation_funcs = [bit_flip, swap_mutate, inversion_mutate]


    for crossfunc in crossover_funcs:
        for mutatefunc in mutation_funcs:
            solution, time_taken, fitness = genetic_algorithm(result,
                 crossfunc,mutatefunc, weights, values, capacity, population_size, generations, mutation_rate
            )
            results[ crossfunc.__name__ +" "+ mutatefunc.__name__] = (solution, time_taken, fitness)  
    return results




### ---------EXAMPLE---------

#20 items
weights = [59, 8, 72, 29, 53, 67, 39, 62, 73, 69, 17, 59, 82, 72, 57, 81, 94, 33, 79, 72]
values = [22, 2, 86, 93, 84, 42, 6, 15, 79, 56, 19, 10, 14, 43, 83, 69, 39, 75, 68, 45]
capacity = 250

# # 10 items
# weights = [84, 73, 47, 19, 36, 39, 14, 28, 4, 53]
# values = [84, 73, 98, 70, 68, 43, 32, 87, 64, 31]
# capacity = 150




max_value, selected_items = bruteforce_knapsack(weights, values, capacity)

data = run_all_combinations(max_value,weights, values, capacity)




from docx import Document
from docx.shared import Inches

# Dữ liệu từ các bảng
data = [
    {
        "Combination": ["one_point_crossover bit_flip", "one_point_crossover swap_mutate", "one_point_crossover inversion_mutate", 
                        "two_points_crossover bit_flip", "two_points_crossover swap_mutate", "two_points_crossover inversion_mutate",
                        "uniform_crossover bit_flip", "uniform_crossover swap_mutate", "uniform_crossover inversion_mutate"],
        "ValueEfficiency": [1.0, 1.0, 1.0, 0.8622, 1.0, 0.8551, 1.0, 1.0, 1.0],
        "Generation": [17.0, 47.0, 28.0, 99.0, 98.0, 99.0, 77.0, 18.0, 33.0],
        "Time Taken": [0.1232, 0.3856, 0.2300, 0.8310, 0.7856, 0.7121, 0.5613, 0.1388, 0.2293]
    },
    # ... (Dữ liệu của 4 bảng còn lại)
]

# Tạo tài liệu Word
document = Document()

# Thêm từng bảng vào tài liệu
for table_data in data:
    table = document.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Combination'
    hdr_cells[1].text = 'ValueEfficiency'
    hdr_cells[2].text = 'Generation'
    hdr_cells[3].text = 'Time Taken'

    for combination, value_efficiency, generation, time_taken in zip(
        table_data["Combination"], table_data["ValueEfficiency"], table_data["Generation"], table_data["Time Taken"]
    ):
        row_cells = table.add_row().cells
        row_cells[0].text = combination
        row_cells[1].text = str(value_efficiency)
        row_cells[2].text = str(generation)
        row_cells[3].text = str(time_taken)
    
    # Thêm một đoạn trống giữa các bảng
    document.add_paragraph("")

# Lưu tài liệu Word
document.save('knapsack_results.docx')